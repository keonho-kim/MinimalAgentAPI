from pathlib import Path
from typing import Any


def apply_docx_edit(
    *,
    path: Path,
    operation: str,
    slots: dict[str, str],
    source_filename: str,
) -> list[dict[str, Any]]:
    from docx import Document

    document = Document(path)
    changed_count = 0
    if operation == "replace_text":
        old_text = slots.get("OLD_TEXT")
        new_value = slots.get("NEW_TEXT")
        if old_text is None or new_value is None:
            raise ValueError("DOCX replace_text requires OLD_TEXT and NEW_TEXT.")
        for paragraph in document.paragraphs:
            if old_text in paragraph.text:
                paragraph.text = paragraph.text.replace(old_text, new_value)
                changed_count += 1
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if old_text in cell.text:
                        cell.text = cell.text.replace(old_text, new_value)
                        changed_count += 1
    elif operation == "add_paragraph":
        new_value = slots.get("TEXT")
        if new_value is None:
            raise ValueError("DOCX add_paragraph requires TEXT.")
        document.add_paragraph(new_value)
        changed_count = 1
    elif operation == "update_table_cell":
        table_index = int(slots.get("TABLE", "1")) - 1
        row_index = int(slots.get("ROW", "1")) - 1
        column_index = int(slots.get("COLUMN", "1")) - 1
        new_value = slots.get("TEXT")
        if new_value is None:
            raise ValueError("DOCX update_table_cell requires TEXT.")
        document.tables[table_index].cell(row_index, column_index).text = new_value
        changed_count = 1
    else:
        raise ValueError(f"Unsupported DOCX edit operation: {operation}")
    if changed_count == 0:
        raise ValueError("DOCX edit did not change the document.")
    document.save(path)
    return [
        {
            "source_file": source_filename,
            "operation": operation,
            "new_value": new_value,
            "changed_count": changed_count,
        }
    ]
